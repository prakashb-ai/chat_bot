"""
import streamlit as st
import pandas as pd
import docx


# Predefined file path
file_path = "/home/prakash/Documents/finrapt company/Azure-Cognitive-Search-Azure-OpenAI-Accelerator/apps/frontend/streamlit_project/9781509302963_Microsoft Azure Essentials Fundamentals of Azure 2nd ed mobile.docx"

# Load the Excel file into a DataFrame
def load_data(file_path):
    try:
        if not file_path.endswith('.docx'):
            st.error("Unsupported file format. Please provide a Word (.docx) file.")
            return None
        
        doc = docx.Document(file_path)
        data = {'Text': []}
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                data['Text'].append(text)
        
        if not data['Text']:
            st.error("No text found in the document.")
            return None
        
        return pd.DataFrame(data)
    
    except FileNotFoundError:
        st.error("File not found. Please provide a valid file path.")
        return None
    
    except Exception as e:
        st.error(f"An error occurred: {e}")
        return None

# Function to search for answer based on user input
def search_question(doc_data, question):
    try:
        if question.strip():
            words = question.lower().split()
            matching_paragraphs = []
            
            for word in words:
                matches = [text for text in doc_data['Text'] if word in text.lower()]
                matching_paragraphs.extend(matches)
            
            if matching_paragraphs:
                related_text = ' '.join(matching_paragraphs)
                related_text_words = related_text.split()[:min(500, len(related_text.split()))]
                related_text = ' '.join(related_text_words)
                
                if len(related_text_words) >= 5:
                    return related_text
                else:
                    return "Not enough related content found (less than 50 words)."
            else:
                return "Word not found in document."
        else:
            return "Please provide a valid question."

    except Exception as e:
        return f"An error occurred while searching for the answer: {e}"

    
api_key = "sk-proj-mStZJGxKY7a1Q8E7PuOuT3BlbkFJIwayTEmzvcFO7tDKNanx"


def main():
    st.title("Simple Chatbot")
    st.markdown("Welcome to the simple chatbot")

    # Load the data from the Word document
    qa_data = load_data(file_path)
    
    if qa_data is not None:
        # Input text box for user input
        user_input = st.text_input("You: write only words", "", max_chars=300)

        # Button to send the user input
        if st.button("Send", key="send_button"):
            if user_input.strip() != "":
                # Get and display bot response
                bot_response = search_question(qa_data, user_input)
                # Style for bot response box
                bot_response_style = """
                  #  width: 600px; 
                   # height: 200px; 
                  #  overflow-y: auto; 
                    #border: 2px solid #ccc; 
                  #  padding: 10px; 
                   # margin-top: 10px; 
                   # border-radius: 5px;
                   # background-color: #f9f9f9;
                
              #  st.markdown(f'<div style="{bot_response_style}"><strong>Bot:</strong> {bot_response}</div>', unsafe_allow_html=True)

#if __name__ == "__main__":
 #   main()
"""


"""

import streamlit as st
import pandas as pd
import docx

# Predefined file path
file_path = "/home/prakash/Documents/finrapt company/Azure-Cognitive-Search-Azure-OpenAI-Accelerator/apps/frontend/streamlit_project/9781509302963_Microsoft Azure Essentials Fundamentals of Azure 2nd ed mobile.docx"

# Load the Word document and extract text, headings, and paragraphs
# Load the Word document and extract text, headings, and paragraphs
def load_data(file_path):
    try:
        if not file_path.endswith('.docx'):
            st.error("Unsupported file format. Please provide a Word (.docx) file.")
            return None
        
        doc = docx.Document(file_path)
        data = {'Heading': [], 'Content': [], 'AllText': []}
        
        current_heading = None
        current_content = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                data['AllText'].append(text)
                if any(run.bold for run in paragraph.runs if run.bold is not None):
                    if current_heading:
                        # Save the current heading and its content
                        data['Heading'].append(current_heading)
                        data['Content'].append(' '.join(current_content))
                    # Start a new heading
                    current_heading = text
                    current_content = []
                else:
                    if current_heading:
                        current_content.append(text)
                    else:
                        # If no heading detected, consider the paragraph as content
                        current_content.append(text)
        
        if current_heading:
            # Save the last heading and its content
            data['Heading'].append(current_heading)
            data['Content'].append(' '.join(current_content))
        
        # Ensure all lists have the same length
        max_length = max(len(data['Heading']), len(data['Content']), len(data['AllText']))
        data['Heading'] += [''] * (max_length - len(data['Heading']))
        data['Content'] += [''] * (max_length - len(data['Content']))
        data['AllText'] += [''] * (max_length - len(data['AllText']))
        
        # Debugging: Print lengths of the lists
      
        
        if not data['Heading']:
            st.error("No headings found in the document.")
            return None
        
        return pd.DataFrame(data)
    
    except FileNotFoundError:
        st.error("File not found. Please provide a valid file path.")
        return None
    
    except Exception as e:
        st.error(f"An error occurred: {e}")
        return None

# Function to search for content based on the provided heading or keyword
def search_question(doc_data, keyword):
    try:
        if keyword.strip():
            keyword = keyword.lower()

            # Search for partial heading matches
            heading_matches = [f"**{heading}**\n\n{content}" for heading, content in zip(doc_data['Heading'], doc_data['Content']) if keyword in heading.lower()]
            
            if heading_matches:
                related_text = '\n\n'.join(heading_matches)
                return related_text

            # If no heading match, search all text for keyword
            text_matches = [text for text in doc_data['AllText'] if keyword in text.lower()]
            
            if text_matches:
                related_text = ' '.join(text_matches)
                return related_text
            else:
                return "Keyword not found in document."
        else:
            return "Please provide a valid keyword."

    except Exception as e:
        return f"An error occurred while searching for the keyword: {e}"

def check_word_in_document(sentence, doc_data):
    try:
        # Tokenize the sentence into words
        words = sentence.lower().split()

        # Iterate through the words in the sentence
        for word in words:
            # Check if the word is present in any of the document headings
            for heading, content in zip(doc_data['Heading'], doc_data['Content']):
                if word.lower() in heading.lower():
                    # Return the related content
                    return content
        # If no match is found, return a message
        return "No relevant content found in the document."
    except Exception as e:
        return f"An error occurred while searching for the word: {e}"


def main():
    st.title("Simple Chatbot")
    st.markdown("Welcome to the simple chatbot. Type a heading or a keyword and press Enter to get related content.")
    with st.sidebar:
        st.markdown("""# Instructions""")
        markdown_content = """
        Ask a question only headings 
        
        Chapter 1: Getting started with Microsoft \n
            * Azure 
            * What is Azure?  
            * Overview of cloud computing
            * Cloud offering
            * Azure services 
            * The new world: Azure Resource Manager  
            * What is it?
            * Why use Resource Manager?
            * Maximize the benefits of using Resource
            * Manager  
            * Resource group tips  
            * Tips for using Resource Manager templates 
            * The classic deployment model
            * PowerShell changes for the Resource Manager
            * and classic deployment models 
            * Role-Based Access Control
            * What is it?
            * Roles
            * Custom roles 
            * The Azure portal
            * Dashboard and hub 
            * Creating and viewing resources
            * Subscription management and billing  
            * Available subscriptions  
            * Share administrative privileges for your
            * Azure subscription 
            * Pricing calculator 
            * Viewing billing in the Azure portal
            * Azure Billing API
            * Azure documentation and samples
            * Documentation
            * Samples 


            Chapter 2: Azure App Service and Web Apps  \n
            * App Service and App Service plans
            * What is an App Service? 
            * So what is an App Service plan?
            * How does this help you? 
            * How to create an App Service plan in the Azure portal 
            * Creating and deploying Web Apps
            * What is a Web App? 
            * Options for creating Web Apps
            * Demo: Create a web app by using the Azure Marketplace 
            * Demo: Create an ASP.NET website in Visual
            * Studio and deploy it as a web app 
            * Configuring, scaling, and monitoring Web Apps 
            * Configuring Web Apps  
            * Monitoring Web Apps
            * Scaling Web Apps 

            Chapter 3: Azure Virtual Machines   \n
            * What is Azure Virtual Machines?
            * Billing
            * Service level agreement 
            * Virtual machine models
            * Azure Resource Manager model
            * Classic/Azure Service Management model 
            * Virtual machine components 
            * Virtual machine 
            * Disks
            * Virtual Network 
            * Availability set
            * Create virtual machines 
            * Create a virtual machine with the Azure portal 
            * Create a virtual machine with a template 
            * Connecting to a virtual machine
            * Remotely access a virtual machine 
            * Network connectivity
            * Configuring and managing a virtual machine
            * Disks
            * Fault domains and update domains 
            * Image capture 
            * Scaling Azure Virtual Machines  
            * Resource Manager virtual machines 
            * Classic virtual machines 

            Chapter 4: Azure Storage  \n
            * Storage accounts  
            * General-purpose storage accounts
            * Blob storage accounts 
            * Storage services
            * Blob storage 
            * File storage 
            * Table storage 
            * Queue storage 
            * Redundancy  
            * Security and Azure Storage 
            * Securing your storage account
            * Securing access to your data
            * Securing your data in transit
            * Encryption at rest 
            * Using Storage Analytics to audit access 
            * Using Cross-Origin Resource Sharing (CORS)  
            * Creating and managing storage
            * Create a storage account using the Azure portal 
            * Create a container and upload blobs using
            * Visual Studio Cloud Explorer
            * Create a file share and upload files using the Azure portal 
            * Create a table and add records using the Visual Studio Cloud Explorer 
            * Create a storage account using PowerShell  
            * Create a container and upload blobs using PowerShell 
            * Create a file share and upload files using PowerShell 
            * AzCopy: A very useful tool 
            * The Azure Data Movement Library 

            Chapter 5: Azure Virtual Networks   \n
            * What is a virtual network (VNet)?
            * Overview 
            * Definitions
            * Creating a virtual network 
            * Creating a virtual network using the Azure portal
            * Creating a virtual network using a Resource
            * Manager template 
            * Network Security Groups  
            * Cross-premises connection options 
            * Site-to-site connectivity 
            * Point-to-site connectivity 
            * Comparing site-to-site and point-to-site connectivity 
            * Private site-to-site connectivity (ExpressRoute) 
            * Point-to-site network 
            * Overview of setup process 
            * Configuring point-to-site VPN

            Chapter 6: Databases   \n
            * Azure SQL Database 
            * Administration 
            * Billing 
            * Business continuity 
            * Applications connecting to SQL Database 
            * SQL Server in Azure Virtual Machines 
            * Billing 
            * Virtual machine configuration
            * Business continuity 
            * Comparing SQL Database with SQL Server in Azure Virtual Machines 
            * Database alternatives 
            * MySQL 
            * NoSQL options 

            Chapter 7: Azure Active Directory  \n
            * Overview of Azure Active Directory 
            * What is Azure Active Directory?
            * Active Directory editions 
            * Creating a directory 
            * Custom domains 
            * Delete a directory 
            * Users and groups  
            * Add users
            * Add groups 
            * Azure Multi-Factor Authentication 
            * Application gallery 
            * Adding gallery applications .
            * Assigning users to application
            * MyApps 

            Chapter 8: Management tools  \n
            * Management tools overview 
            * Visual Studio 2015 and the Azure SDK 
            * Install the Azure SDK 
            * Manage resources with Cloud Explorer 
            * Create an Azure resource Windows PowerShell  
            * Azure PowerShell cmdlet installation 
            * Connecting to Azure 
            * Cross-platform command-line interface  
            * Installation 
            * Connecting to Azure 
            * Usage 

            Chapter 9: Additional Azure services   \n 
            * Some other Azure services we think you
            * should know about  
            * Azure Service Fabric 
            * Cloud Services 
            * Azure Container Service  
            * DocumentDB 
            * Azure Redis Cache  
            * Azure HDInsight 
            * Azure Search 
            * Azure Service Bus 
            * Azure Event Hubs 
            * Azure Notification Hubs  
            * Azure Media Services 
            * Azure Backup 
            * Azure Site Recovery 
            * Azure Key Vault 
            * More Azure services  

            Chapter 10: Business cases   \n
            * Development and test scenarios  
            * Hybrid scenarios  
            * Network connectivity 
            * Internet connectivity 
            * Application and infrastructure modernization and migration  
            * Azure Mobile Apps  
            * Machine learning   

        """

        st.markdown(markdown_content)

    # Load the data from the Word document
    qa_data = load_data(file_path)
    
    if qa_data is not None:
        # Input text box for user input
        user_input = st.text_input("You: (type a heading or keyword)", "", max_chars=100)

        # Button to send the user input
        if st.button("Send", key="send_button"):
            if user_input.strip() != "":
                # Get and display bot response
                bot_response = search_question(qa_data, user_input)
                
                # Style for bot response box
                bot_response_style = """
                    width: 600px; 
                    height: 400px; 
                    overflow-y: auto; 
                    border: 2px solid #ccc; 
                    padding: 10px; 
                    margin-top: 10px; 
                    border-radius: 5px;
                    background-color: #f9f9f9;
                """
                st.markdown(f'<div style="{bot_response_style}"><strong>Bot:</strong><br>{bot_response}</div>', unsafe_allow_html=True)

if __name__ == "__main__":
    main()
